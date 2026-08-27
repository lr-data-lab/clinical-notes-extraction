"""Exact-match evaluation for the medication extraction pipeline.

Mirrors ``schemas.py``. Any change there must be applied here: the attribute
list below is the set of scored slots, and a key present in one file and absent
from the other is silently unscored.

Scoring unit
------------
One *slot* = one key of one annotation: ``flag_is_medication_completed``, plus
the ten attributes of every medication.

    TP  the annotation holds a value and the prediction reproduces it exactly
    FN  the annotation holds a value and the prediction does not
    FP  the prediction holds a value the annotation does not

Absent on both sides is a true negative and is not counted: most attributes are
filled only occasionally, so rewarding agreement on emptiness would give every
run a high score for leaving keys blank. A slot present on both sides but
unequal counts as one FP *and* one FN, so precision and recall diverge only
through omissions and inventions.

Matching criterion
------------------
Equality is exact after the canonicalisation in :func:`canonical`, applied
symmetrically to both sides: ``null`` and ``[]`` both mean absent, booleans
become ``"true"``/``"false"`` so ``false`` is a value, whitespace runs are
collapsed, and the three array fields are compared as ordered tuples because the
template requires their elements in source order. Nothing else is normalised:
the annotation rule is verbatim, so casing and punctuation are part of the
value. Collapsing whitespace is the one departure from that rule, and is safe
here because the template forbids newlines in every scored field; it forgives
only a difference in spacing, which no reviewer would call an extraction error.

Failed extractions
------------------
A note whose extraction failed, or which the run never produced, is scored as
*no output*: every annotated slot becomes a false negative and no false
positive is created. Recall and F1 absorb the failures; precision does not.
``invalid`` reports their share and must be read next to the scores.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Iterable, Mapping, NamedTuple, Sequence

import numpy as np
import pandas as pd
from scipy.optimize import linear_sum_assignment

# --------------------------------------------------------------------------- #
# Schema — keep in sync with schemas.py
# --------------------------------------------------------------------------- #

#: MedicationAttributes, in schema order
ATTRIBUTES: tuple[str, ...] = (
    "active_substance",
    "commercial_name",
    "dosage_form",
    "dose",
    "quantity",
    "route",
    "frequency",
    "duration",
    "indication",
    "administration_instructions",
)

#: Scored top-level keys. ``medications_text`` is excluded from the slot-level
#: micro-average because it is a single span whose exact match is binary and
#: would otherwise weigh the same as one attribute of one medication. It is a
#: genuine extraction decision -- the template requires finding the heading and
#: truncating before any inpatient/discharge sub-section -- so it is reported
#: separately by :func:`section_table` rather than left unmeasured.
TOP_LEVEL_KEYS: tuple[str, ...] = ("flag_is_medication_completed",)

SCORED_KEYS: tuple[str, ...] = TOP_LEVEL_KEYS + ATTRIBUTES

#: Arrays compared order-INsensitively. Empty by design: the template's rule 5
#: requires ``dose``, ``frequency`` and ``indication`` to hold "one element per
#: distinct value stated, IN SOURCE ORDER", so order is part of the annotation
#: rule for all three and a reordered array is a genuine error, not a formatting
#: artefact. Adding a key here is a deliberate leniency and must be stated in the
#: Methods as a departure from the annotation rule.
SET_KEYS: frozenset[str] = frozenset()

#: minimum token overlap for two medications to be treated as the same drug
MATCH_THRESHOLD: float = 0.20

_TOKEN = re.compile(r"[a-z0-9]+")
_PUNCT = re.compile(r"[^a-z0-9]+")


# --------------------------------------------------------------------------- #
# Comparison
# --------------------------------------------------------------------------- #


def canonical(key: str, value: Any) -> Any:
    """Canonical form of one value; ``None`` means absent."""
    if value is None:
        return None
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (list, tuple)):
        items = [c for c in (canonical(key, v) for v in value) if c is not None]
        if not items:
            return None
        return frozenset(items) if key in SET_KEYS else tuple(items)
    text = " ".join(str(value).split())
    return text or None


def _counts(gold: Any, pred: Any) -> tuple[int, int, int]:
    if gold is None and pred is None:
        return 0, 0, 0
    if gold == pred:
        return 1, 0, 0
    return 0, int(pred is not None), int(gold is not None)


def _category(gold: Any, pred: Any) -> str:
    """How a slot is wrong, not only that it is."""

    def flat(value: Any) -> str | None:
        if value is None:
            return None
        if isinstance(value, (frozenset, tuple)):
            return " | ".join(sorted(map(str, value)))
        return str(value)

    g, p = flat(gold), flat(pred)
    if g is None:
        return "hallucinated"
    if p is None:
        return "missed"
    if g.casefold() == p.casefold():
        return "casing"
    gs, ps = _PUNCT.sub("", g.casefold()), _PUNCT.sub("", p.casefold())
    if gs and ps and (gs in ps or ps in gs):
        return "span_boundary"
    return "different_value"


# --------------------------------------------------------------------------- #
# Aligning the two medication lists
# --------------------------------------------------------------------------- #


def _tokens(medication: Mapping[str, Any]) -> set[str]:
    attributes = medication.get("attributes") or {}
    text = " ".join(
        str(value)
        for value in (
            attributes.get("active_substance"),
            attributes.get("commercial_name"),
            medication.get("span_text"),
        )
        if value
    )
    return set(_TOKEN.findall(text.casefold()))


def align(
    gold_meds: Sequence[Mapping[str, Any]],
    pred_meds: Sequence[Mapping[str, Any]],
    threshold: float = MATCH_THRESHOLD,
) -> tuple[list[tuple[int, int]], list[int], list[int], list[float]]:
    """Match medications on identity before their attributes are compared.

    A position-by-position comparison would let one missed medication cascade an
    error onto every attribute of every medication after it, which measures list
    alignment rather than extraction quality. Similarity is the Jaccard overlap
    of the tokens of the two name fields and the source span — all three are
    present on both sides, so the measure is symmetric — and the assignment
    maximising the total is found with the Hungarian algorithm.
    """
    if not gold_meds or not pred_meds:
        return [], list(range(len(gold_meds))), list(range(len(pred_meds))), []

    gold_tokens = [_tokens(m) for m in gold_meds]
    pred_tokens = [_tokens(m) for m in pred_meds]
    sim = np.array(
        [[len(g & p) / len(g | p) if g and p else 0.0 for p in pred_tokens] for g in gold_tokens]
    )

    rows, cols = linear_sum_assignment(-sim)
    pairs = [(int(i), int(j)) for i, j in zip(rows, cols) if sim[i, j] >= threshold]
    matched_g = {i for i, _ in pairs}
    matched_p = {j for _, j in pairs}
    return (
        pairs,
        [i for i in range(len(gold_meds)) if i not in matched_g],
        [j for j in range(len(pred_meds)) if j not in matched_p],
        [float(sim[i, j]) for i, j in pairs],
    )


# --------------------------------------------------------------------------- #
# Scoring
# --------------------------------------------------------------------------- #


class Scored(NamedTuple):
    """Three tidy frames; every table below is a groupby on these."""

    slots: pd.DataFrame          # model, strategy, note_id, key, tp, fp, fn
    notes: pd.DataFrame          # model, strategy, note_id, valid, medication counts
    disagreements: pd.DataFrame  # one row per slot where the two sides differ


def usable_output(record: Mapping[str, Any] | None) -> Mapping[str, Any] | None:
    """The prediction, or ``None`` if the record is not usable.

    A record flagged ``status="ok"`` but carrying a null output is a failure and
    is counted as one.
    """
    if record is None or record.get("status", "ok") != "ok":
        return None
    output = record.get("output")
    return output if isinstance(output, Mapping) else None


def score(
    gold_by_id: Mapping[str, Mapping[str, Any]],
    records: Iterable[Mapping[str, Any]],
    *,
    threshold: float = MATCH_THRESHOLD,
) -> Scored:
    """Score every (model, strategy) over every annotated note."""
    predictions: dict[tuple[str, str], dict[str, Any]] = {}
    for record in records:
        cell = (str(record.get("model", "")), str(record.get("strategy", "")))
        note_id = str(record.get("note_id", ""))
        # a usable output always wins over an error file for the same note, so
        # the result never depends on the order the directories were walked
        if predictions.setdefault(cell, {}).get(note_id) is None:
            predictions[cell][note_id] = usable_output(record)

    slots: list[dict[str, Any]] = []
    notes: list[dict[str, Any]] = []
    disagreements: list[dict[str, Any]] = []

    for cell in sorted(predictions):
        model, strategy = cell
        for note_id, gold in gold_by_id.items():
            prediction = predictions[cell].get(note_id)
            valid = prediction is not None
            prediction = prediction or {}
            tag = {"model": model, "strategy": strategy, "note_id": note_id}

            def add(key: str, gold_value: Any, pred_value: Any, medication: str) -> None:
                g, p = canonical(key, gold_value), canonical(key, pred_value)
                tp, fp, fn = _counts(g, p)
                if not (tp or fp or fn):
                    return
                slots.append({**tag, "key": key, "tp": tp, "fp": fp, "fn": fn})
                if not tp:
                    disagreements.append({
                        **tag, "medication": medication, "key": key,
                        "gold": gold_value, "predicted": pred_value,
                        "category": _category(g, p), "valid": valid,
                    })

            for key in TOP_LEVEL_KEYS:
                add(key, gold.get(key), prediction.get(key), "-")

            gold_meds = gold.get("medications") or []
            pred_meds = prediction.get("medications") or []
            pairs, only_gold, only_pred, similarities = align(gold_meds, pred_meds, threshold)

            for i, j in pairs:
                g_attrs = gold_meds[i].get("attributes") or {}
                p_attrs = pred_meds[j].get("attributes") or {}
                for key in ATTRIBUTES:
                    add(key, g_attrs.get(key), p_attrs.get(key), f"g{i + 1}~p{j + 1}")
            for i in only_gold:
                g_attrs = gold_meds[i].get("attributes") or {}
                for key in ATTRIBUTES:
                    add(key, g_attrs.get(key), None, f"g{i + 1}~-")
            for j in only_pred:
                p_attrs = pred_meds[j].get("attributes") or {}
                for key in ATTRIBUTES:
                    add(key, None, p_attrs.get(key), f"-~p{j + 1}")

            section_gold = canonical("medications_text", gold.get("medications_text"))
            section_pred = canonical("medications_text", prediction.get("medications_text"))

            notes.append({
                **tag,
                "valid": valid,
                "section_gold_present": section_gold is not None,
                "section_exact": section_gold == section_pred,
                "n_gold_medications": len(gold_meds),
                "n_pred_medications": len(pred_meds),
                "med_tp": len(pairs), "med_fp": len(only_pred), "med_fn": len(only_gold),
                "min_similarity": min(similarities) if similarities else np.nan,
            })

    return Scored(
        pd.DataFrame(slots, columns=["model", "strategy", "note_id", "key", "tp", "fp", "fn"]),
        pd.DataFrame(notes),
        pd.DataFrame(disagreements),
    )


# --------------------------------------------------------------------------- #
# Tables
# --------------------------------------------------------------------------- #


def prf(frame: pd.DataFrame) -> pd.DataFrame:
    """Add P, R, F1 and support to a frame of summed tp/fp/fn.

    Precision is ``NaN`` where nothing was predicted: undefined, not a measured
    zero. F1 is ``2TP / (2TP + FP + FN)``, which equals the harmonic mean of P
    and R wherever both are defined and stays defined when one is not.
    """
    tp, fp, fn = frame["tp"], frame["fp"], frame["fn"]
    out = frame.copy()
    out["P"] = np.where(tp + fp > 0, tp / (tp + fp), np.nan)
    out["R"] = np.where(tp + fn > 0, tp / (tp + fn), np.nan)
    out["F1"] = np.where(2 * tp + fp + fn > 0, 2 * tp / (2 * tp + fp + fn), np.nan)
    out["support"] = tp + fn
    return out


def headline_table(scored: Scored) -> pd.DataFrame:
    """One row per cell: P, R, F1, invalid, support."""
    totals = scored.slots.groupby(["model", "strategy"])[["tp", "fp", "fn"]].sum()
    table = prf(totals)
    table["invalid"] = 1 - scored.notes.groupby(["model", "strategy"])["valid"].mean()
    return table[["P", "R", "F1", "invalid", "support"]].sort_values("F1", ascending=False)


def medication_table(scored: Scored) -> pd.DataFrame:
    """Detection of the medication objects themselves.

    Every attribute score is conditional on detection, so this belongs in the
    error analysis rather than the headline. Missed detections are already
    inside the attribute micro-average as false negatives, so demoting this
    table under-reports nothing.
    """
    totals = (
        scored.notes.groupby(["model", "strategy"])[["med_tp", "med_fp", "med_fn"]]
        .sum()
        .rename(columns={"med_tp": "tp", "med_fp": "fp", "med_fn": "fn"})
    )
    return prf(totals)[["P", "R", "F1", "support"]]


def section_table(scored: Scored) -> pd.DataFrame:
    """Exact match on ``medications_text``, the section span itself.

    Reported apart from the slot micro-average: it is one span per note, so
    pooling it with the attribute slots would let a whole section weigh the same
    as a single ``route``. ``exact`` is the proportion of notes whose section
    span is reproduced exactly (whitespace collapsed); ``n_section`` is the
    number of notes where the annotation holds a section at all.
    """
    notes = scored.notes
    table = notes.groupby(["model", "strategy"]).agg(
        exact=("section_exact", "mean"),
        n_notes=("section_exact", "size"),
        n_section=("section_gold_present", "sum"),
    )
    return table.sort_values("exact", ascending=False)


def per_key_table(scored: Scored, model: str | None = None) -> pd.DataFrame:
    """One row per key, P/R/F1 per cell, plus Macro and Micro.

    Keys with no annotated value anywhere are dropped: an all-zero row describes
    the split, not the model. ``Macro`` is the unweighted mean over the rows
    shown, so a rare attribute weighs as much as ``active_substance``; ``Micro``
    pools their counts and reproduces the headline F1.
    """
    slots = scored.slots if model is None else scored.slots[scored.slots.model == model]

    wide = (
        prf(slots.groupby(["model", "strategy", "key"])[["tp", "fp", "fn"]].sum())
        .reset_index()
        .pivot(index="key", columns=["model", "strategy"], values=["P", "R", "F1"])
        .reorder_levels([1, 2, 0], axis=1)
    )
    keys = [k for k in SCORED_KEYS if k in wide.index]
    cells = sorted({(c[0], c[1]) for c in wide.columns})
    wide = wide.reindex(index=keys, columns=pd.MultiIndex.from_tuples(
        [(m, s, metric) for m, s in cells for metric in ("P", "R", "F1")]
    ))

    totals = prf(slots.groupby(["model", "strategy"])[["tp", "fp", "fn"]].sum())
    micro = pd.Series({c: totals.loc[(c[0], c[1]), c[2]] for c in wide.columns}, name="Micro")
    table = pd.concat([wide, wide.mean(axis=0).rename("Macro").to_frame().T, micro.to_frame().T])

    support = slots.groupby("key")[["tp", "fn"]].sum().sum(axis=1).reindex(keys)
    table.insert(0, "support", list(support) + [support.sum(), support.sum()])

    if model is not None:  # one model per table: the model level is constant
        table.columns = pd.MultiIndex.from_tuples(
            [("", c) if isinstance(c, str) else (c[1], c[2]) for c in table.columns]
        )
    table.index.name = "Key"
    return table


# --------------------------------------------------------------------------- #
# Uncertainty
# --------------------------------------------------------------------------- #


def note_matrices(scored: Scored, note_ids: Sequence[str]) -> dict[tuple[str, str], np.ndarray]:
    """(n_notes, 3) array of tp/fp/fn per note, per cell, in ``note_ids`` order."""
    per_note = scored.slots.groupby(["model", "strategy", "note_id"])[["tp", "fp", "fn"]].sum()
    return {
        cell: frame.droplevel([0, 1]).reindex(note_ids).fillna(0).to_numpy(float)
        for cell, frame in per_note.groupby(level=[0, 1])
    }


def bootstrap_f1(matrix: np.ndarray, boot_index: np.ndarray) -> np.ndarray:
    """F1 recomputed on each resample of the *notes* — the level at which the
    observations are independent. Counts are pooled inside a resample before F1
    is recomputed."""
    sums = matrix[boot_index].sum(axis=1)
    tp, fp, fn = sums[:, 0], sums[:, 1], sums[:, 2]
    denominator = 2 * tp + fp + fn
    return np.divide(2 * tp, denominator, out=np.zeros_like(tp), where=denominator > 0)


def _f1(matrix: np.ndarray) -> float:
    tp, fp, fn = matrix.sum(axis=0)
    denominator = 2 * tp + fp + fn
    return 2 * tp / denominator if denominator else np.nan


def paired_bootstrap(
    matrices: Mapping[tuple[str, str], np.ndarray],
    a: tuple[str, str],
    b: tuple[str, str],
    boot_index: np.ndarray,
) -> dict[str, Any]:
    """Compare two cells on the same notes and the same resamples.

    ``p_two_sided`` uses the ``(r + 1) / (n + 1)`` correction: with a finite
    number of resamples a p-value of exactly zero is not attainable, so
    ``p_is_floor`` marks the rows where it is only an upper bound. This is a
    bootstrap p-value, descriptive rather than a formal test.
    """
    diff = bootstrap_f1(matrices[a], boot_index) - bootstrap_f1(matrices[b], boot_index)
    lo, hi = float(np.percentile(diff, 2.5)), float(np.percentile(diff, 97.5))
    tail = min(int((diff <= 0).sum()), int((diff >= 0).sum()))
    return {
        "a": f"{a[0]}/{a[1]}",
        "b": f"{b[0]}/{b[1]}",
        "delta_F1": _f1(matrices[a]) - _f1(matrices[b]),
        "lo": lo,
        "hi": hi,
        "p_two_sided": min(2 * (tail + 1) / (len(diff) + 1), 1.0),
        "p_is_floor": tail == 0,
        "separates": lo > 0 or hi < 0,
    }


# --------------------------------------------------------------------------- #
# Loading
# --------------------------------------------------------------------------- #


def load_ground_truth(directory: str | Path) -> dict[str, dict[str, Any]]:
    """One annotation per file; the id is ``note_id`` or the file name."""
    notes = {}
    for path in sorted(Path(directory).glob("*.json")):
        note = json.loads(path.read_text(encoding="utf-8"))
        notes[str(note.get("note_id") or path.stem)] = note
    return notes


def load_results(run_dir: str | Path) -> list[dict[str, Any]]:
    """Load ``<run>/<model>/<strategy>/*.json``.

    Failed extractions under ``<run>/<model>/errors/<strategy>/`` are loaded too
    and marked ``status="error"``, so they are scored as missing output instead
    of disappearing from the denominator.
    """
    records = []
    for model_dir in sorted(p for p in Path(run_dir).iterdir() if p.is_dir()):
        for child in sorted(p for p in model_dir.iterdir() if p.is_dir()):
            if child.name == "prompts":
                continue
            failed = child.name == "errors"
            directories = sorted(p for p in child.iterdir() if p.is_dir()) if failed else [child]
            for directory in directories:
                for path in sorted(directory.glob("*.json")):
                    payload = json.loads(path.read_text(encoding="utf-8"))
                    record = (
                        dict(payload)
                        if isinstance(payload, dict) and "output" in payload
                        else {"output": payload}
                    )
                    record.setdefault("note_id", path.stem)
                    record["model"] = record.get("model") or model_dir.name
                    record["strategy"] = record.get("strategy") or directory.name
                    record["status"] = "error" if failed else record.get("status", "ok")
                    record["path"] = str(path)
                    records.append(record)
    return records


# --------------------------------------------------------------------------- #
# Word export
# --------------------------------------------------------------------------- #


def to_docx(tables: Mapping[str, pd.DataFrame], path: str | Path, *, title: str = "") -> Path:
    """Write tables to a .docx, ready to paste into the thesis.

    MultiIndex columns are flattened to ``"strategy F1"``; the index is written
    as the leading column, so set a meaningful index before calling.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    if title:
        document.add_heading(title, level=1)

    for caption, frame in tables.items():
        frame = frame.reset_index()
        headers = [
            " ".join(str(part) for part in column if part != "") if isinstance(column, tuple)
            else str(column)
            for column in frame.columns
        ]

        document.add_paragraph(caption).runs[0].bold = True
        table = document.add_table(rows=len(frame) + 1, cols=len(headers))
        table.style = "Table Grid"

        for j, header in enumerate(headers):
            table.cell(0, j).text = header
        for i, row in enumerate(frame.itertuples(index=False), start=1):
            for j, value in enumerate(row):
                if isinstance(value, float):
                    text = "--" if np.isnan(value) else f"{value:.3f}"
                else:
                    text = str(value)
                table.cell(i, j).text = text

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.bold = i == 0
        document.add_paragraph()

    path = Path(path)
    document.save(path)
    return path